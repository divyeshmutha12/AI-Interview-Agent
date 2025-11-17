"""
AI Interviewer System + Knowledge Base Builder (Merged Streamlit UI)
With KB Rebuild functionality integrated
"""

import streamlit as st
import os
import sys
import tempfile
from pathlib import Path
import logging
import json
import io
from dotenv import load_dotenv

# Add project root to path
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

# ---- IMPORT SERVICES ----
from ai_interviewer.agents.agent_service import AgentService
from ai_interviewer.utils.document_parser import DocumentParser

# Import KB components
from langchain_community.vectorstores import FAISS
from langchain_community.document_loaders import PyMuPDFLoader
from docx import Document as DocxDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain_core.documents import Document as LCDocument

# Import hybrid retriever for rebuild functionality
from ai_interviewer.kb.hybrid_retriever import build_kb, load_kb, IN_MEMORY_KB

# Load environment variables
load_dotenv()

# Setup logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# ✅ Set base directory for storing FAISS indexes
BASE_KB_DIR = Path("stored_kbs")
BASE_KB_DIR.mkdir(exist_ok=True)

# Streamlit Page Config
st.set_page_config(
    page_title="AI Interviewer & KB System",
    page_icon="🎯",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for better UI
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        font-weight: bold;
        text-align: center;
        padding: 1rem 0;
        background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        margin-bottom: 1rem;
    }
    .admin-section {
        background-color: #f8f9fa;
        padding: 1rem;
        border-radius: 0.5rem;
        border-left: 4px solid #dc3545;
        margin: 1rem 0;
    }
    .success-box {
        padding: 1rem;
        border-radius: 0.5rem;
        background-color: #d4edda;
        border: 1px solid #c3e6cb;
        color: #155724;
        margin: 1rem 0;
    }
    .info-box {
        padding: 1rem;
        border-radius: 0.5rem;
        background-color: #d1ecf1;
        border: 1px solid #bee5eb;
        color: #0c5460;
        margin: 1rem 0;
    }
    .warning-box {
        padding: 1rem;
        border-radius: 0.5rem;
        background-color: #fff3cd;
        border: 1px solid #ffeaa7;
        color: #856404;
        margin: 1rem 0;
    }
    .stButton>button {
        border-radius: 0.5rem;
        padding: 0.5rem 1rem;
    }
</style>
""", unsafe_allow_html=True)

# -------------------- KB REBUILD FUNCTIONS --------------------

def needs_bm25_rebuild(kb_name: str) -> bool:
    """Check if a KB needs BM25 rebuild (missing BM25 files)"""
    kb_path = BASE_KB_DIR / kb_name
    has_bm25 = (kb_path / "bm25_tokens.json").exists()
    has_chunks = (kb_path / "chunks.json").exists()
    return not (has_bm25 and has_chunks)

def rebuild_kb_with_bm25(kb_name: str) -> tuple[bool, str]:
    """
    Rebuild an existing KB to add BM25 support.
    Integrated from rebuild_kb.py
    """
    kb_path = BASE_KB_DIR / kb_name

    if not kb_path.exists():
        return False, f"KB '{kb_name}' not found at {kb_path}"

    try:
        # Load existing FAISS index
        embeddings = OpenAIEmbeddings(model="text-embedding-3-small")
        faiss_index = FAISS.load_local(
            str(kb_path),
            embeddings,
            index_name="faiss_index",
            allow_dangerous_deserialization=True
        )

        # Get all documents from FAISS
        docs = list(faiss_index.docstore._dict.values())

        if not docs:
            return False, f"No documents found in KB '{kb_name}'"

        # Extract text from documents
        kb_docs = [doc.page_content for doc in docs]

        # Rebuild with hybrid retriever
        build_kb(kb_name, kb_docs)

        return True, f"✅ Successfully rebuilt KB '{kb_name}' with BM25 support! ({len(docs)} documents)"

    except Exception as e:
        logger.error(f"Error rebuilding KB: {e}")
        return False, f"❌ Error rebuilding KB: {str(e)}"

def get_kb_status(kb_info: dict) -> dict:
    """Get detailed status of a knowledge base"""
    kb_path = BASE_KB_DIR / kb_info["name"]

    status = {
        "name": kb_info["name"],
        "display_name": kb_info["display_name"],
        "role": kb_info["role"],
        "chunk_count": kb_info["chunk_count"],
        "has_faiss": (kb_path / "faiss_index").exists(),
        "has_bm25": (kb_path / "bm25_tokens.json").exists(),
        "has_chunks": (kb_path / "chunks.json").exists(),
    }

    # Determine status
    if status["has_faiss"] and status["has_bm25"] and status["has_chunks"]:
        status["status"] = "✅ Hybrid (FAISS + BM25)"
        status["needs_rebuild"] = False
    elif status["has_faiss"] and not (status["has_bm25"] and status["has_chunks"]):
        status["status"] = "⚠️ FAISS Only (Needs BM25)"
        status["needs_rebuild"] = True
    else:
        status["status"] = "❌ Incomplete"
        status["needs_rebuild"] = True

    return status

# -------------------- KB HELPER FUNCTIONS --------------------

def read_docx_content(file_bytes: bytes) -> str:
    """Extract text content from DOCX bytes."""
    try:
        doc = DocxDocument(io.BytesIO(file_bytes))
        return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    except Exception as e:
        logger.error(f"Error reading DOCX: {e}")
        return ""

def load_uploaded_file(file) -> str:
    """Return extracted text from uploaded PDF or DOCX file."""
    try:
        if file.name.lower().endswith(".pdf"):
            # Temporarily write to buffer for PyMuPDFLoader
            with io.BytesIO(file.getbuffer()) as pdf_bytes:
                temp_path = f"/tmp/{file.name}"
                with open(temp_path, "wb") as temp_file:
                    temp_file.write(pdf_bytes.getvalue())
                loader = PyMuPDFLoader(temp_path)
                docs = loader.load()
                os.remove(temp_path)
                return "\n\n".join([doc.page_content for doc in docs])
        elif file.name.lower().endswith(".docx"):
            text = read_docx_content(file.getvalue())
            return text
        else:
            return file.getvalue().decode("utf-8")
    except Exception as e:
        logger.error(f"Error loading uploaded file: {e}")
        return ""

def create_kb(kb_name: str, role: str, uploaded_files: list):
    """Create FAISS vector store and metadata for uploaded documents."""
    try:
        kb_dir = BASE_KB_DIR / kb_name.replace(" ", "_").lower()
        kb_dir.mkdir(exist_ok=True)

        docs = []
        for file in uploaded_files:
            file_content = load_uploaded_file(file)
            if file_content.strip():
                docs.append(LCDocument(page_content=file_content, metadata={"source": file.name}))

        if not docs:
            return False, "No valid content found in uploaded files"

        # Split into chunks
        splitter = RecursiveCharacterTextSplitter(chunk_size=800, chunk_overlap=200)
        chunks = splitter.split_documents(docs)

        # Create FAISS index
        embeddings = OpenAIEmbeddings(model="text-embedding-3-small")
        vector_store = FAISS.from_documents(chunks, embeddings)

        # Save index locally
        vector_store.save_local(str(kb_dir / "faiss_index"))

        # Save metadata
        metadata = {"name": kb_name, "role": role, "chunk_count": len(chunks)}
        with open(kb_dir / "metadata.json", "w") as f:
            json.dump(metadata, f)

        # ✅ ALSO BUILD HYBRID KB using the new function
        kb_docs = [doc.page_content for doc in docs]
        build_kb(kb_name.replace(" ", "_").lower(), kb_docs)

        return True, f"Knowledge Base '{kb_name}' created successfully with {len(chunks)} chunks! (Hybrid: FAISS + BM25)"

    except Exception as e:
        logger.error(f"Error creating KB: {e}")
        return False, f"Error creating knowledge base: {str(e)}"

def search_kb(kb_name: str, query: str, k: int = 3):
    """Search within an existing KB."""
    try:
        kb_dir = BASE_KB_DIR / kb_name
        embeddings = OpenAIEmbeddings(model="text-embedding-3-small")

        vector_store = FAISS.load_local(
            str(kb_dir / "faiss_index"),
            embeddings,
            allow_dangerous_deserialization=True
        )
        return vector_store.similarity_search(query, k=k)
    except Exception as e:
        logger.error(f"Error searching KB: {e}")
        return []

def get_available_kbs():
    """Get list of available knowledge bases with status."""
    kbs = []
    if BASE_KB_DIR.exists():
        for item in BASE_KB_DIR.iterdir():
            if item.is_dir() and (item / "faiss_index").exists():
                # Load metadata
                metadata_path = item / "metadata.json"
                if metadata_path.exists():
                    with open(metadata_path, 'r') as f:
                        metadata = json.load(f)
                    kb_info = {
                        "name": item.name,
                        "display_name": metadata.get("name", item.name),
                        "role": metadata.get("role", "Unknown"),
                        "chunk_count": metadata.get("chunk_count", 0)
                    }
                    # Get detailed status
                    kbs.append(get_kb_status(kb_info))
                else:
                    kb_info = {
                        "name": item.name,
                        "display_name": item.name,
                        "role": "Unknown",
                        "chunk_count": 0
                    }
                    kbs.append(get_kb_status(kb_info))
    return kbs

# -------------------- MAIN APP FUNCTIONS --------------------

@st.cache_resource
def initialize_services():
    """Initialize agent service and document parser (cached)"""
    try:
        agent_service = AgentService()
        doc_parser = DocumentParser()
        return agent_service, doc_parser, None
    except Exception as e:
        logger.error(f"Failed to initialize services: {e}")
        return None, None, str(e)

def save_uploaded_file(uploaded_file) -> str:
    """Save uploaded file to temporary location"""
    try:
        temp_dir = Path(tempfile.gettempdir()) / "ai_interviewer_uploads"
        temp_dir.mkdir(exist_ok=True)
        file_path = temp_dir / uploaded_file.name
        with open(file_path, "wb") as f:
            f.write(uploaded_file.getbuffer())
        return str(file_path)
    except Exception as e:
        logger.error(f"Error saving file: {e}")
        raise

def main():
    """Main application"""

    # Header
    st.markdown('<h1 class="main-header">🎯 AI Interviewer & Knowledge Base System</h1>', unsafe_allow_html=True)
    st.markdown(
        '<p class="sub-header">Generate Interview Questions • Evaluate Answers • Build Knowledge Bases</p>',
        unsafe_allow_html=True
    )

    # Initialize services
    with st.spinner("Initializing AI agents and services..."):
        agent_service, doc_parser, init_error = initialize_services()

    if init_error:
        st.error(f"❌ Initialization Failed: {init_error}")
        st.info("Please check your .env file and ensure all API keys are set correctly.")
        st.stop()

    # Sidebar
    with st.sidebar:
        st.header("⚙️ Configuration")

        # Interview settings
        num_questions = st.slider(
            "Number of Questions",
            min_value=1,
            max_value=20,
            value=5,
            help="How many interview questions to generate"
        )

        # Model info
        model = os.getenv("OPENAI_MODEL", "gpt-4o-mini")
        st.info(f"**Model:** {model}")

        # Supported formats
        with st.expander("📄 Supported Formats"):
            formats = doc_parser.get_supported_formats()
            st.write(", ".join(formats))

        st.markdown("---")
        st.markdown("### 📊 Monitoring")
        st.markdown("[View Traces in Langfuse](https://cloud.langfuse.com)")

    # Main tabs - MERGED: Interview + KB functionality + Admin
    tab1, tab2, tab3, tab4 = st.tabs([
        "📤 Interview Questions",
        "✅ Evaluate Answers",
        "📚 Knowledge Bases",
        "⚙️ Admin Tools"
    ])

    # ===== TAB 1: Interview Questions =====
    with tab1:
        st.header("🎯 Generate Interview Questions")

        col1, col2 = st.columns(2)

        with col1:
            st.subheader("📄 Candidate Resume")
            resume_file = st.file_uploader(
                "Upload Resume",
                type=['pdf', 'docx', 'txt'],
                help="Upload candidate's resume",
                key="resume_upload"
            )
            if resume_file:
                st.success(f"✅ {resume_file.name}")

        with col2:
            st.subheader("💼 Job Description")
            jd_file = st.file_uploader(
                "Upload Job Description",
                type=['pdf', 'docx', 'txt'],
                help="Upload job description",
                key="jd_upload"
            )
            if jd_file:
                st.success(f"✅ {jd_file.name}")

        # KB Integration for Question Generation
        st.markdown("---")
        st.subheader("📚 Knowledge Base Integration (Optional)")

        available_kbs = get_available_kbs()
        kb_names = [kb["display_name"] for kb in available_kbs]

        if kb_names:
            selected_kb = st.selectbox(
                "Use Existing Knowledge Base",
                options=["None"] + kb_names,
                help="Select a pre-built KB to enhance question generation"
            )
        else:
            selected_kb = "None"
            st.info("No knowledge bases available. Create one in the Knowledge Bases tab.")

        # Generate button
        if st.button("🚀 Generate Interview Questions", type="primary", use_container_width=True):
            if not resume_file or not jd_file:
                st.error("⚠️ Please upload both Resume and Job Description")
            else:
                try:
                    with st.spinner("Processing documents..."):
                        # Save and parse uploaded files
                        resume_path = save_uploaded_file(resume_file)
                        jd_path = save_uploaded_file(jd_file)

                        docs = doc_parser.parse_resume_and_jd(resume_path, jd_path)
                        resume_text = docs['resume']
                        jd_text = docs['job_description']

                    # Prepare KB context
                    kb_name_to_use = None
                    if selected_kb != "None":
                        # Find the actual KB name from display name
                        for kb in available_kbs:
                            if kb["display_name"] == selected_kb:
                                kb_name_to_use = kb["name"]
                                break

                    # Generate questions with KB context
                    with st.spinner("🤖 AI agents are generating personalized interview questions..."):
                        result = agent_service.generate_interview_questions(
                            candidate_resume=resume_text,
                            job_profile=jd_text,
                            num_questions=num_questions,
                            session_id=None
                        )

                    # Display results
                    if result.get('error'):
                        st.error(f"❌ Error: {result['error']}")
                    else:
                        st.success("✅ Interview questions generated successfully!")

                        # Show KB usage info
                        if kb_name_to_use:
                            st.info(f"🧠 Enhanced with Knowledge Base: **{selected_kb}**")

                        # Candidate Analysis (if available)
                        if result.get('candidate_analysis'):
                            with st.expander("🔍 Candidate Analysis", expanded=False):
                                st.markdown(result['candidate_analysis'])

                        # Questions
                        st.markdown("---")
                        st.subheader("📋 Generated Interview Questions")
                        st.markdown(result['result'])

                        # Download option
                        st.download_button(
                            label="📥 Download Questions",
                            data=result['result'],
                            file_name="interview_questions.txt",
                            mime="text/plain",
                            use_container_width=True
                        )

                except Exception as e:
                    st.error(f"❌ Error processing documents: {str(e)}")
                    logger.error(f"Processing error: {e}")

    # ===== TAB 2: Evaluate Answers =====
    with tab2:
        st.header("✅ Evaluate Candidate Answers")

        st.markdown("Enter a question and the candidate's answer below to get an AI-powered evaluation.")

        question_text = st.text_area(
            "Interview Question",
            height=100,
            placeholder="e.g., What is a closure in Python?",
            help="Enter the interview question that was asked"
        )

        answer_text = st.text_area(
            "Candidate's Answer",
            height=200,
            placeholder="Enter the candidate's response here...",
            help="Paste or type the candidate's answer"
        )

        if st.button("🔍 Evaluate Answer", type="primary", use_container_width=True):
            if not question_text or not answer_text:
                st.error("⚠️ Please provide both question and answer")
            else:
                try:
                    eval_request = f"Question: {question_text}\n\nAnswer: {answer_text}"

                    with st.spinner("🤖 Evaluating answer..."):
                        result = agent_service.run(
                            user_input=eval_request,
                            session_id=None
                        )

                    if result.get('error'):
                        st.error(f"❌ Error: {result['error']}")
                    else:
                        st.success("✅ Evaluation complete!")
                        st.markdown("---")
                        st.subheader("📊 Evaluation Results")
                        st.markdown(result['result'])

                except Exception as e:
                    st.error(f"❌ Error evaluating answer: {str(e)}")

    # ===== TAB 3: Knowledge Bases =====
    with tab3:
        st.header("📚 Knowledge Base Management")

        kb_tab1, kb_tab2 = st.tabs(["➕ Create New KB", "🔍 Search Existing KBs"])

        # ---- Create KB Tab ----
        with kb_tab1:
            st.subheader("Create a New Knowledge Base")

            col1, col2 = st.columns(2)

            with col1:
                kb_name = st.text_input(
                    "Knowledge Base Name",
                    placeholder="e.g., Python Backend KB",
                    help="Give your knowledge base a descriptive name"
                )

            with col2:
                role_tag = st.text_input(
                    "Role/Technology Tag",
                    placeholder="e.g., python_backend, data_science, web_dev",
                    help="Tag for categorizing this KB"
                )

            st.subheader("Upload Documents")
            uploaded_docs = st.file_uploader(
                "Choose PDF, DOCX, or TXT files",
                type=["pdf", "docx", "txt"],
                accept_multiple_files=True,
                help="Upload documents to build your knowledge base"
            )

            # Show uploaded files
            if uploaded_docs:
                st.write(f"📁 **Uploaded Files ({len(uploaded_docs)}):**")
                for doc in uploaded_docs:
                    st.write(f"• {doc.name} ({doc.size / 1024:.1f} KB)")

            if st.button("🏗️ Create Knowledge Base", type="primary", use_container_width=True):
                if not kb_name or not role_tag or not uploaded_docs:
                    st.warning("⚠️ Please fill all fields and upload at least one document.")
                else:
                    with st.spinner("Building knowledge base..."):
                        success, message = create_kb(kb_name, role_tag, uploaded_docs)

                    if success:
                        st.success(f"✅ {message}")
                    else:
                        st.error(f"❌ {message}")

        # ---- Search KB Tab ----
        with kb_tab2:
            st.subheader("Search Existing Knowledge Bases")

            available_kbs = get_available_kbs()

            if not available_kbs:
                st.info("ℹ️ No knowledge bases found. Create one in the 'Create New KB' tab first.")
            else:
                # KB Selection
                kb_options = {kb["display_name"]: kb for kb in available_kbs}
                selected_kb_display = st.selectbox(
                    "Select Knowledge Base",
                    options=list(kb_options.keys()),
                    help="Choose which knowledge base to search"
                )

                selected_kb = kb_options[selected_kb_display]

                # Show KB info with status
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    st.metric("KB Name", selected_kb["display_name"])
                with col2:
                    st.metric("Role Tag", selected_kb["role"])
                with col3:
                    st.metric("Chunks", selected_kb["chunk_count"])
                with col4:
                    st.metric("Status", "Hybrid" if not selected_kb["needs_rebuild"] else "FAISS Only")

                # Search interface
                st.subheader("🔍 Search Query")
                search_query = st.text_input(
                    "Enter your search query",
                    placeholder="e.g., Python decorators, system design patterns...",
                    help="Search for specific topics or concepts in the knowledge base"
                )

                num_results = st.slider("Number of results", 1, 10, 3)

                if st.button("🔎 Search Knowledge Base", type="primary", use_container_width=True):
                    if not search_query:
                        st.warning("⚠️ Please enter a search query.")
                    else:
                        with st.spinner(f"Searching '{selected_kb_display}'..."):
                            results = search_kb(selected_kb["name"], search_query, k=num_results)

                        if not results:
                            st.info("🤷 No results found for your query. Try different keywords.")
                        else:
                            st.success(f"✅ Found {len(results)} results")
                            st.markdown("---")

                            for i, result in enumerate(results, 1):
                                with st.expander(f"📄 Result {i} - {result.metadata.get('source', 'Document')}", expanded=i==1):
                                    st.write(result.page_content)
                                    st.caption(f"Source: {result.metadata.get('source', 'Unknown')}")

    # ===== TAB 4: Admin Tools (Rebuild KB Functionality) =====
    with tab4:
        st.header("⚙️ Admin Tools")

        st.markdown("""
        <div class="admin-section">
        <h3>🔧 Knowledge Base Maintenance</h3>
        <p>Tools for maintaining and upgrading your knowledge bases.</p>
        </div>
        """, unsafe_allow_html=True)

        st.subheader("🔄 Rebuild Knowledge Bases")
        st.info("""
        **Purpose:** Convert older FAISS-only knowledge bases to hybrid format (FAISS + BM25).

        **When to use:**
        - If you have KBs created before hybrid search was implemented
        - To add BM25 keyword search capability to existing KBs
        - KBs showing 'FAISS Only' status in the Knowledge Bases tab
        """)

        available_kbs = get_available_kbs()
        kbs_needing_rebuild = [kb for kb in available_kbs if kb["needs_rebuild"]]

        if not available_kbs:
            st.warning("No knowledge bases found. Create some KBs first.")
        elif not kbs_needing_rebuild:
            st.success("🎉 All knowledge bases are already in hybrid format (FAISS + BM25)!")

            # Show current KB status
            st.subheader("📊 Current KB Status")
            for kb in available_kbs:
                col1, col2, col3, col4 = st.columns([2, 2, 1, 1])
                with col1:
                    st.write(f"**{kb['display_name']}**")
                with col2:
                    st.write(kb['role'])
                with col3:
                    st.write(f"{kb['chunk_count']} chunks")
                with col4:
                    st.success("✅ Hybrid")
        else:
            st.warning(f"Found {len(kbs_needing_rebuild)} knowledge bases that need BM25 rebuild:")

            for kb in kbs_needing_rebuild:
                with st.expander(f"🔧 {kb['display_name']} - {kb['status']}", expanded=True):
                    col1, col2 = st.columns([3, 1])

                    with col1:
                        st.write(f"**Role:** {kb['role']}")
                        st.write(f"**Chunks:** {kb['chunk_count']}")
                        st.write(f"**Current:** FAISS only")
                        st.write(f"**Target:** FAISS + BM25 hybrid")

                    with col2:
                        if st.button(f"Rebuild", key=f"rebuild_{kb['name']}"):
                            with st.spinner(f"Rebuilding {kb['display_name']}..."):
                                success, message = rebuild_kb_with_bm25(kb['name'])

                            if success:
                                st.success(message)
                                st.rerun()  # Refresh the UI
                            else:
                                st.error(message)

            # Bulk rebuild option
            st.markdown("---")
            st.subheader("🚀 Bulk Rebuild All")

            if st.button("🔄 Rebuild All Incomplete KBs", type="primary", use_container_width=True):
                success_count = 0
                error_count = 0

                progress_bar = st.progress(0)
                status_text = st.empty()

                for i, kb in enumerate(kbs_needing_rebuild):
                    status_text.text(f"Rebuilding {kb['display_name']}... ({i+1}/{len(kbs_needing_rebuild)})")

                    success, message = rebuild_kb_with_bm25(kb['name'])
                    if success:
                        success_count += 1
                    else:
                        error_count += 1
                        st.error(f"Failed to rebuild {kb['display_name']}: {message}")

                    progress_bar.progress((i + 1) / len(kbs_needing_rebuild))

                status_text.text("✅ Rebuild complete!")
                st.success(f"Rebuilt {success_count} KBs successfully, {error_count} failures")
                st.rerun()

        # Additional admin tools can be added here
        st.markdown("---")
        st.subheader("📈 System Information")

        col1, col2 = st.columns(2)
        with col1:
            st.metric("Total Knowledge Bases", len(available_kbs))
            st.metric("Hybrid KBs", len([kb for kb in available_kbs if not kb["needs_rebuild"]]))

        with col2:
            st.metric("FAISS-Only KBs", len(kbs_needing_rebuild))
            st.metric("Total Chunks", sum(kb["chunk_count"] for kb in available_kbs))

if __name__ == "__main__":
    main()